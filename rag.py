from google import genai
import chromadb
from chromadb.utils import embedding_functions
from docx import Document
from rank_bm25 import BM25Okapi
import os

# ── API key ───────────────────────────────────────────
import os

API_KEY = os.getenv("GEMINI_API_KEY")

# ── Daftar dokumen ────────────────────────────────────
DOKUMEN = [
    "S01_Ide_AI_Assistant_Bob.docx",
    "S02_Prompt_and_Safety_Plan_Bob.docx",
    "S03_Simulasi_Alur_Bob.docx",
    "S04_Hizkia Reinaldy_24110400036_Agent_Workflow_Design.docx",
    "S06_Hizkia Reinaldy_24110400036_Retrieval_Semantic_Search_Plan.docx",
]


# ── Baca teks dari .docx (paragraf + tabel) ───────────
def baca_docx(nama_file):
    doc = Document(nama_file)
    teks = ""
    # Baca paragraf
    for para in doc.paragraphs:
        if para.text.strip():
            teks += para.text.strip() + "\n"
    # Baca tabel
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    teks += cell.text.strip() + "\n"
    return teks


# ── Chunking dengan overlap ───────────────────────────
def buat_chunks(teks, nama_file, ukuran=300, overlap=50):
    kata = teks.split()
    chunks = []
    i = 0
    idx = 0
    while i < len(kata):
        chunk = " ".join(kata[i : i + ukuran])
        chunks.append(
            {"id": f"{nama_file}_chunk_{idx}", "text": chunk, "source": nama_file}
        )
        i += ukuran - overlap
        idx += 1
    return chunks


# ── Query Expansion ───────────────────────────────────
def expand_query(pertanyaan, client):
    prompt = f"""Berikan 3 variasi pertanyaan berbeda dari pertanyaan berikut untuk memperluas pencarian.
Format: hanya tulis 3 pertanyaan, satu per baris, tanpa nomor atau penjelasan tambahan.
Pertanyaan: {pertanyaan}"""
    response = client.models.generate_content(model="gemini-3.5-flash", contents=prompt)
    variasi = response.text.strip().split("\n")
    semua = [pertanyaan] + [v.strip() for v in variasi if v.strip()]
    return semua[:4]


# ── Context Compression ───────────────────────────────
def compress_context(konteks, pertanyaan, client):
    prompt = f"""Dari konteks berikut, ambil HANYA bagian yang relevan dengan pertanyaan.
Buang informasi yang tidak relevan. Pertahankan informasi penting.

Pertanyaan: {pertanyaan}

Konteks:
{konteks}

Tulis ulang hanya bagian yang relevan, maksimal 500 kata."""
    response = client.models.generate_content(model="gemini-3.5-flash", contents=prompt)
    return response.text.strip()


# ── Setup ChromaDB ────────────────────────────────────
print("Menyiapkan database RAG...")
chroma_client = chromadb.Client()
embedding_fn = embedding_functions.SentenceTransformerEmbeddingFunction(
    model_name="all-MiniLM-L6-v2"
)
collection = chroma_client.get_or_create_collection(
    name="bob_ai_docs", embedding_function=embedding_fn
)

# ── Load dokumen ──────────────────────────────────────
print("Memuat dokumen ke database...")
semua_chunks = []
for nama_file in DOKUMEN:
    if os.path.exists(nama_file):
        print(f"  ✓ Membaca: {nama_file}")
        teks = baca_docx(nama_file)
        chunks = buat_chunks(teks, nama_file)
        semua_chunks.extend(chunks)
    else:
        print(f"  ✗ Tidak ditemukan: {nama_file}")

if semua_chunks:
    collection.add(
        ids=[c["id"] for c in semua_chunks],
        documents=[c["text"] for c in semua_chunks],
        metadatas=[{"source": c["source"]} for c in semua_chunks],
    )
    print(f"\n✅ {len(semua_chunks)} chunks berhasil dimuat!\n")


# ── Fungsi RAG lengkap ────────────────────────────────
def tanya_rag(pertanyaan, tampilkan_proses=True):
    client = genai.Client(api_key=API_KEY)

    # 1. Query Expansion
    if tampilkan_proses:
        print("\n[1] Query Expansion...")
    queries = expand_query(pertanyaan, client)
    if tampilkan_proses:
        for q in queries:
            print(f"    → {q}")

    # 2. Retrieval — ambil chunks relevan dari semua query
    if tampilkan_proses:
        print("\n[2] Retrieval...")
    semua_hasil = {}
    for q in queries:
        hasil = collection.query(query_texts=[q], n_results=3)
        for i, doc in enumerate(hasil["documents"][0]):
            chunk_id = hasil["ids"][0][i]
            if chunk_id not in semua_hasil:
                semua_hasil[chunk_id] = {
                    "text": doc,
                    "source": hasil["metadatas"][0][i]["source"],
                    "score": 0,
                }

    # 3. Reranking dengan BM25
    if tampilkan_proses:
        print("\n[3] Reranking dengan BM25...")
    corpus = [v["text"].split() for v in semua_hasil.values()]
    bm25 = BM25Okapi(corpus)
    skor = bm25.get_scores(pertanyaan.split())
    chunks_list = list(semua_hasil.values())
    for i, chunk in enumerate(chunks_list):
        chunk["score"] = skor[i]
    chunks_sorted = sorted(chunks_list, key=lambda x: x["score"], reverse=True)[:3]

    if tampilkan_proses:
        print(f"    → {len(chunks_sorted)} chunks terpilih setelah reranking")

    # 4. Tampilkan Retrieved Context
    konteks_raw = ""
    sumber_list = []
    print("\n" + "=" * 55)
    print("📄 RETRIEVED CONTEXT:")
    print("=" * 55)
    for i, chunk in enumerate(chunks_sorted):
        print(f"\n[Chunk {i+1}] Sumber: {chunk['source']}")
        print(f"Score BM25: {chunk['score']:.4f}")
        print(f"Isi: {chunk['text'][:200]}...")
        konteks_raw += f"\n---\n{chunk['text']}"
        if chunk["source"] not in sumber_list:
            sumber_list.append(chunk["source"])

    # 5. Context Compression
    if tampilkan_proses:
        print("\n[4] Context Compression...")
    konteks_compressed = compress_context(konteks_raw, pertanyaan, client)

    # 6. Generate jawaban dengan LLM
    if tampilkan_proses:
        print("\n[5] Generating jawaban dengan Gemini...")
    prompt_final = f"""Kamu adalah Bob AI, asisten berbahasa Indonesia untuk mahasiswa.
Jawab pertanyaan berikut HANYA berdasarkan konteks dokumen yang diberikan.
Jika jawaban tidak ada dalam konteks, katakan dengan jujur bahwa informasi tersebut tidak tersedia dalam dokumen.

Konteks:
{konteks_compressed}

Pertanyaan: {pertanyaan}

Berikan jawaban yang jelas, terstruktur, dan mudah dipahami dalam bahasa Indonesia.
Akhiri dengan: 'Periksa sumber resminya ya untuk memastikan keakuratan informasi ini.'"""

    response = client.models.generate_content(
        model="gemini-3.5-flash", contents=prompt_final
    )

    return response.text, sumber_list


# ── Loop percakapan ───────────────────────────────────
print("=" * 55)
print("  Bob AI — RAG Mode (Advanced)")
print("  Query Expansion + Reranking + Compression")
print("  Ketik 'keluar' untuk berhenti")
print("=" * 55)

while True:
    pertanyaan = input("\nKamu  : ")
    if pertanyaan.strip().lower() == "keluar":
        print("\nBob AI: Sampai jumpa! Semangat demo UTS-nya ya! 🎉")
        break
    if pertanyaan.strip() == "":
        continue

    print("\nBob AI sedang memproses...")
    jawaban, sumber = tanya_rag(pertanyaan)

    print("\n" + "=" * 55)
    print("💬 JAWABAN BOB AI:")
    print("=" * 55)
    print(jawaban)
    print("\n📚 Sumber dokumen:")
    for s in sumber:
        print(f"   • {s}")
    print("=" * 55)
