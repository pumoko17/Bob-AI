"""Bob AI - Integrated Application

A comprehensive Streamlit application combining simple chat, RAG mode,
and dataset management for Indonesian language productivity tasks.
"""

import os
from io import BytesIO

import streamlit as st
from google import genai
import pdfplumber
from docx import Document
import chromadb
from chromadb.utils import embedding_functions
from rank_bm25 import BM25Okapi
import kagglehub
import pandas as pd

# API key dari environment variable
API_KEY = os.getenv("GEMINI_API_KEY")

# System prompt
SYSTEM_PROMPT = """Kamu adalah Bob AI, asisten produktivitas berbahasa Indonesia.
Tugasmu membantu pengguna meringkas dokumen, membuat jadwal belajar,
memperbaiki tulisan, dan menjawab pertanyaan sehari-hari.
Jawab dengan bahasa Indonesia yang jelas, singkat, dan mudah dipahami.
Selalu akhiri jawaban dengan kalimat:
'Periksa sumber resminya ya untuk memastikan keakuratan informasi ini.'"""

# Konfigurasi halaman
st.set_page_config(page_title="Bob AI - Integrated", page_icon="🤖", layout="wide")

# Inisialisasi session state
if "messages" not in st.session_state:
    st.session_state.messages = []
if "riwayat_gemini" not in st.session_state:
    st.session_state.riwayat_gemini = []
if "pdf_dimuat" not in st.session_state:
    st.session_state.pdf_dimuat = False
if "docx_dimuat" not in st.session_state:
    st.session_state.docx_dimuat = False
if "rag_mode" not in st.session_state:
    st.session_state.rag_mode = False
if "chroma_collection" not in st.session_state:
    st.session_state.chroma_collection = None
if "semua_chunks" not in st.session_state:
    st.session_state.semua_chunks = []

# Sidebar
with st.sidebar:
    st.header("🤖 Bob AI Settings")

    # Mode selection
    st.subheader("Mode Operasi")
    mode = st.radio(
        "Pilih Mode:",
        ["Simple Chat", "RAG Mode (Advanced)", "Dataset Mode"],
        label_visibility="collapsed",
    )

    if mode == "RAG Mode (Advanced)":
        st.session_state.rag_mode = True
        st.info("📚 RAG Mode: Query Expansion + Reranking + Context Compression")
    else:
        st.session_state.rag_mode = False

    st.divider()

    # Document Upload
    st.subheader("📄 Upload Dokumen")

    uploaded_pdf = st.file_uploader("Upload PDF", type=["pdf"], key="pdf_upload")
    uploaded_docx = st.file_uploader("Upload DOCX", type=["docx"], key="docx_upload")

    # Process PDF
    if uploaded_pdf and not st.session_state.pdf_dimuat:
        with st.spinner("Membaca PDF..."):
            try:
                with pdfplumber.open(uploaded_pdf) as pdf:
                    teks = ""
                    for halaman in pdf.pages:
                        teks += halaman.extract_text() or ""
                if teks.strip():
                    PESAN_PDF = (
                        f"Berikut isi dokumen '{uploaded_pdf.name}':\n\n"
                        f"{teks[:3000]}\n\n"
                        "Konfirmasi bahwa kamu sudah membaca dokumen ini."
                    )
                    st.session_state.riwayat_gemini.append(
                        {"role": "user", "parts": [{"text": PESAN_PDF}]}
                    )
                    client = genai.Client(api_key=API_KEY)
                    response = client.models.generate_content(
                        model="gemini-3.5-flash",
                        contents=st.session_state.riwayat_gemini,
                        config={"system_instruction": SYSTEM_PROMPT},
                    )
                    konfirmasi = response.text
                    st.session_state.riwayat_gemini.append(
                        {"role": "model", "parts": [{"text": konfirmasi}]}
                    )
                    st.session_state.messages.append(
                        {
                            "role": "assistant",
                            "content": f"📄 **{uploaded_pdf.name}** berhasil dibaca!\n\n{konfirmasi}",
                        }
                    )
                    st.session_state.pdf_dimuat = True
                    st.success("PDF berhasil dimuat!")
            except IOError as e:
                st.error(f"Error reading PDF: {e}")
            except Exception as e:
                st.error(f"Error: {e}")

    # Process DOCX for RAG
    if uploaded_docx and st.session_state.rag_mode:
        with st.spinner("Memproses DOCX untuk RAG..."):
            try:
                doc = Document(BytesIO(uploaded_docx.read()))
                teks = ""
                for para in doc.paragraphs:
                    if para.text.strip():
                        teks += para.text.strip() + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                teks += cell.text.strip() + "\n"

                if teks.strip():
                    # Chunking
                    kata = teks.split()
                    chunks = []
                    i = 0
                    idx = 0
                    while i < len(kata):
                        chunk = " ".join(kata[i : i + 300])
                        chunks.append(
                            {
                                "id": f"{uploaded_docx.name}_chunk_{idx}",
                                "text": chunk,
                                "source": uploaded_docx.name,
                            }
                        )
                        i += 250
                        idx += 1

                    st.session_state.semua_chunks.extend(chunks)
                    st.session_state.docx_dimuat = True
                    st.success(
                        f"✅ {len(chunks)} chunks dari {uploaded_docx.name} berhasil dibuat!"
                    )
            except Exception as e:
                st.error(f"Error processing DOCX: {e}")

    # Initialize ChromaDB for RAG
    if (
        st.session_state.rag_mode
        and st.session_state.semua_chunks
        and not st.session_state.chroma_collection
    ):
        with st.spinner("Menyiapkan RAG database..."):
            try:
                chroma_client = chromadb.Client()
                embedding_fn = embedding_functions.SentenceTransformerEmbeddingFunction(
                    model_name="all-MiniLM-L6-v2"
                )
                collection = chroma_client.get_or_create_collection(
                    name="bob_ai_docs", embedding_function=embedding_fn
                )
                collection.add(
                    ids=[c["id"] for c in st.session_state.semua_chunks],
                    documents=[c["text"] for c in st.session_state.semua_chunks],
                    metadatas=[
                        {"source": c["source"]} for c in st.session_state.semua_chunks
                    ],
                )
                st.session_state.chroma_collection = collection
                st.success(
                    f"✅ {len(st.session_state.semua_chunks)} chunks dimuat ke database!"
                )
            except Exception as e:
                st.error(f"Error initializing RAG: {e}")

    # Status
    st.divider()
    if st.session_state.pdf_dimuat:
        st.success("✅ PDF sudah dimuat")
    if st.session_state.docx_dimuat:
        st.success("✅ DOCX sudah dimuat untuk RAG")
    if st.session_state.chroma_collection:
        st.success("✅ RAG Database aktif")

    # Reset button
    st.divider()
    if st.button("🗑️ Reset Semua"):
        st.session_state.messages = []
        st.session_state.riwayat_gemini = []
        st.session_state.pdf_dimuat = False
        st.session_state.docx_dimuat = False
        st.session_state.semua_chunks = []
        st.session_state.chroma_collection = None
        st.rerun()

# Main content
st.title("🤖 Bob AI - Integrated")
st.caption("Asisten Produktivitas Berbahasa Indonesia dengan RAG")

# Dataset Mode
if mode == "Dataset Mode":
    st.header("📊 Dataset Management")

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("Download & Cek Dataset")
        if st.button("Download World University Rankings"):
            with st.spinner("Downloading dataset..."):
                try:
                    path = kagglehub.dataset_download(
                        "mylesoneill/world-university-rankings"
                    )
                    st.success(f"✅ Dataset downloaded to: {path}")
                    st.write("Files available:")
                    for f in os.listdir(path):
                        st.write(f"  - {f}")
                except Exception as e:
                    st.error(f"Error downloading dataset: {e}")

        if st.button("Download US Education Dataset"):
            with st.spinner("Downloading dataset..."):
                try:
                    path = kagglehub.dataset_download(
                        "noriuk/us-education-datasets-unification-project"
                    )
                    st.success(f"✅ Dataset downloaded to: {path}")
                    st.write("Files available:")
                    for f in os.listdir(path):
                        st.write(f"  - {f}")
                except Exception as e:
                    st.error(f"Error downloading dataset: {e}")

    with col2:
        st.subheader("Siapkan Dataset untuk RAG")
        if st.button("Generate University Rankings Text"):
            with st.spinner("Processing dataset..."):
                try:
                    path = kagglehub.dataset_download(
                        "mylesoneill/world-university-rankings"
                    )
                    df_times = pd.read_csv(os.path.join(path, "timesData.csv"))
                    df_shanghai = pd.read_csv(os.path.join(path, "shanghaiData.csv"))

                    teks_output = []

                    # Process Times Data
                    df_times_clean = df_times.dropna(
                        subset=["university_name", "world_rank"]
                    )
                    df_times_latest = df_times_clean[
                        df_times_clean["year"] == df_times_clean["year"].max()
                    ]
                    for _, row in df_times_latest.head(100).iterrows():
                        teks = (
                            f"Menurut Times Higher Education ranking tahun {row['year']}, "
                            f"{row['university_name']} dari {row['country']} berada di peringkat {row['world_rank']} dunia. "
                            f"Skor pengajaran (teaching): {row['teaching']}, "
                            f"skor riset (research): {row['research']}, "
                            f"skor sitasi (citations): {row['citations']}."
                        )
                        teks_output.append(teks)

                    # Process Shanghai Data
                    df_shanghai_clean = df_shanghai.dropna(
                        subset=["university_name", "world_rank"]
                    )
                    df_shanghai_latest = df_shanghai_clean[
                        df_shanghai_clean["year"] == df_shanghai_clean["year"].max()
                    ]
                    for _, row in df_shanghai_latest.head(100).iterrows():
                        teks = (
                            f"Menurut Shanghai (ARWU) ranking tahun {int(row['year'])}, "
                            f"{row['university_name']} dari {row['country']} berada di peringkat {row['world_rank']} dunia. "
                            f"Total skor alumni: {row['alumni']}, skor publikasi: {row['pub']}, "
                            f"skor penghargaan (award): {row['award']}."
                        )
                        teks_output.append(teks)

                    OUTPUT_FILE = "dataset_university_rankings.txt"
                    with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
                        for teks in teks_output:
                            f.write(teks + "\n\n")

                    st.success(f"✅ {len(teks_output)} entri disimpan ke {OUTPUT_FILE}")
                    st.download_button(
                        label="Download Dataset Text",
                        data=open(OUTPUT_FILE, "r", encoding="utf-8").read(),
                        file_name="dataset_university_rankings.txt",
                        mime="text/plain",
                    )
                except Exception as e:
                    st.error(f"Error processing dataset: {e}")

    st.info(
        "💡 Setelah mendownload dataset, upload file .txt yang dihasilkan ke mode RAG untuk digunakan dalam pencarian."
    )

# Chat Interface (Simple Chat & RAG Mode)
else:
    # Display chat history
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # Chat input
    if prompt := st.chat_input("Ketik pertanyaan kamu di sini..."):
        # Display user message
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        # Process based on mode
        with st.chat_message("assistant"):
            with st.spinner("Bob AI sedang berpikir..."):
                try:
                    client = genai.Client(api_key=API_KEY)

                    if st.session_state.rag_mode and st.session_state.chroma_collection:
                        # RAG Mode with Query Expansion
                        st.info("🔍 Menggunakan RAG dengan Query Expansion...")

                        # Query Expansion
                        EXPAND_PROMPT = (
                            f"Berikan 3 variasi pertanyaan berbeda dari pertanyaan berikut "
                            f"untuk memperluas pencarian.\n"
                            "Format: hanya tulis 3 pertanyaan, satu per baris, "
                            "tanpa nomor atau penjelasan tambahan.\n"
                            f"Pertanyaan: {prompt}"
                        )
                        response = client.models.generate_content(
                            model="gemini-3.5-flash", contents=EXPAND_PROMPT
                        )
                        variasi = response.text.strip().split("\n")
                        queries = [prompt] + [v.strip() for v in variasi if v.strip()]

                        # Retrieval
                        semua_hasil = {}
                        for q in queries[:4]:
                            hasil = st.session_state.chroma_collection.query(
                                query_texts=[q], n_results=3
                            )
                            for i, doc in enumerate(hasil["documents"][0]):
                                chunk_id = hasil["ids"][0][i]
                                if chunk_id not in semua_hasil:
                                    semua_hasil[chunk_id] = {
                                        "text": doc,
                                        "source": hasil["metadatas"][0][i]["source"],
                                        "score": 0,
                                    }

                        # Reranking with BM25
                        corpus = [v["text"].split() for v in semua_hasil.values()]
                        bm25 = BM25Okapi(corpus)
                        skor = bm25.get_scores(prompt.split())
                        chunks_list = list(semua_hasil.values())
                        for i, chunk in enumerate(chunks_list):
                            chunk["score"] = skor[i]
                        chunks_sorted = sorted(
                            chunks_list, key=lambda x: x["score"], reverse=True
                        )[:3]

                        # Context Compression
                        KONTEKS_RAW = "\n---\n".join([c["text"] for c in chunks_sorted])
                        COMPRESS_PROMPT = (
                            f"Dari konteks berikut, ambil HANYA bagian yang relevan "
                            f"dengan pertanyaan.\n"
                            "Buang informasi yang tidak relevan. "
                            "Pertahankan informasi penting.\n\n"
                            f"Pertanyaan: {prompt}\n\n"
                            f"Konteks:\n{KONTEKS_RAW}\n\n"
                            "Tulis ulang hanya bagian yang relevan, maksimal 500 kata."
                        )
                        response = client.models.generate_content(
                            model="gemini-3.5-flash", contents=COMPRESS_PROMPT
                        )
                        konteks_compressed = response.text.strip()

                        # Generate final answer
                        PROMPT_FINAL = (
                            "Kamu adalah Bob AI, asisten berbahasa Indonesia.\n"
                            "Jawab pertanyaan berikut HANYA berdasarkan konteks dokumen "
                            "yang diberikan.\n"
                            "Jika jawaban tidak ada dalam konteks, katakan dengan jujur "
                            "bahwa informasi tersebut tidak tersedia.\n\n"
                            f"Konteks:\n{konteks_compressed}\n\n"
                            f"Pertanyaan: {prompt}\n\n"
                            "Berikan jawaban yang jelas, terstruktur, dan mudah dipahami "
                            "dalam bahasa Indonesia.\n"
                            "Akhiri dengan: 'Periksa sumber resminya ya untuk memastikan "
                            "keakuratan informasi ini.'"
                        )

                        response = client.models.generate_content(
                            model="gemini-3.5-flash", contents=PROMPT_FINAL
                        )
                        jawaban = response.text

                        st.markdown(jawaban)

                        # Show sources
                        sumber_list = list(set([c["source"] for c in chunks_sorted]))
                        st.caption(f"📚 Sumber: {', '.join(sumber_list)}")

                        st.session_state.riwayat_gemini.append(
                            {"role": "user", "parts": [{"text": prompt}]}
                        )
                        st.session_state.riwayat_gemini.append(
                            {"role": "model", "parts": [{"text": jawaban}]}
                        )
                        st.session_state.messages.append(
                            {"role": "assistant", "content": jawaban}
                        )

                    else:
                        # Simple Chat Mode
                        st.session_state.riwayat_gemini.append(
                            {"role": "user", "parts": [{"text": prompt}]}
                        )

                        response = client.models.generate_content(
                            model="gemini-3.5-flash",
                            contents=st.session_state.riwayat_gemini,
                            config={"system_instruction": SYSTEM_PROMPT},
                        )
                        jawaban = response.text
                        st.markdown(jawaban)

                        st.session_state.riwayat_gemini.append(
                            {"role": "model", "parts": [{"text": jawaban}]}
                        )
                        st.session_state.messages.append(
                            {"role": "assistant", "content": jawaban}
                        )

                except genai.errors.APIError as e:
                    st.error(f"API Error: {e}")
                except Exception as e:
                    st.error(f"Error: {e}")
